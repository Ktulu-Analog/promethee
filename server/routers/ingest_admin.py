# ============================================================================
# Prométhée — Assistant IA avancé
# ============================================================================
# Auteur  : Pierre COUGET ktulu.analog@gmail.com
# Licence : GNU Affero General Public License v3.0 (AGPL-3.0)
#           https://www.gnu.org/licenses/agpl-3.0.html
# Année   : 2026
# ----------------------------------------------------------------------------
# Ce fichier fait partie du projet Prométhée.
# Vous pouvez le redistribuer et/ou le modifier selon les termes de la
# licence AGPL-3.0 publiée par la Free Software Foundation.
# ============================================================================

"""
routers/ingest_admin.py — Routes d'ingestion de répertoires (admin only).

Expose les fonctionnalités de scripts/ingest3.py sous forme d'API REST,
réservées aux administrateurs de Prométhée.

Routes
──────
  GET    /admin/ingest/collections          Liste les collections Qdrant
  POST   /admin/ingest/collections          Crée une collection Qdrant
  DELETE /admin/ingest/collections/{name}   Supprime une collection
  GET    /admin/ingest/status               État de Qdrant + config embedding
  POST   /admin/ingest/run                  Lance une ingestion (SSE streaming)
  POST   /admin/ingest/file                 Ingère un fichier uploadé directement
  GET    /admin/ingest/sources              Liste les sources d'une collection
  DELETE /admin/ingest/sources              Supprime une source d'une collection
"""

import asyncio
import json
import logging
import os
import sys
import tempfile
import uuid
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from server.deps import require_admin, get_current_user_config
from core.request_context import set_user_config

_log = logging.getLogger(__name__)
router = APIRouter()

# ── Import des dépendances d'ingestion ────────────────────────────────────────

# Ajouter la racine du projet au path pour importer core/
_project_root = Path(__file__).parent.parent.parent
if str(_project_root) not in sys.path:
    sys.path.insert(0, str(_project_root))

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct
    _QDRANT_OK = True
except ImportError:
    _QDRANT_OK = False
    _log.warning("[ingest_admin] qdrant-client non installé")

try:
    from openai import OpenAI as _OpenAI
    _OPENAI_OK = True
except ImportError:
    _OPENAI_OK = False
    _log.warning("[ingest_admin] openai non installé")

try:
    from core.rag_engine import _chunk_text, _contextual_prefix_batch, _get_embeddings as _rag_get_embeddings
    _RAG_ENGINE_OK = True
except ImportError:
    _RAG_ENGINE_OK = False
    _rag_get_embeddings = None
    _log.warning("[ingest_admin] rag_engine non disponible — chunker/embedder fallback")

# ── Configuration depuis l'environnement ──────────────────────────────────────

QDRANT_URL              = os.getenv("QDRANT_URL", "http://localhost:6333")
QDRANT_API_KEY          = os.getenv("QDRANT_API_KEY", "") or None
EMBEDDING_API_BASE      = os.getenv("EMBEDDING_API_BASE", "https://api.openai.com/v1")
OPENAI_API_KEY          = os.getenv("OPENAI_API_KEY", "")
EMBEDDING_MODEL         = os.getenv("EMBEDDING_MODEL", "text-embedding-ada-002")
EMBEDDING_DIMENSION     = int(os.getenv("EMBEDDING_DIMENSION", "1536"))
QDRANT_BATCH_SIZE       = int(os.getenv("QDRANT_BATCH_SIZE", "100"))
QDRANT_MAX_RETRIES      = int(os.getenv("QDRANT_MAX_RETRIES", "3"))
QDRANT_TIMEOUT          = int(os.getenv("QDRANT_TIMEOUT", "60"))
OCR_TEXT_THRESHOLD      = int(os.getenv("OCR_TEXT_THRESHOLD", "100"))
OCR_LANG                = os.getenv("OCR_LANG", "fra+eng")
OCR_DPI                 = int(os.getenv("OCR_DPI", "300"))
RAG_CONTEXTUAL_CHUNKING = os.getenv("RAG_CONTEXTUAL_CHUNKING", "OFF").strip().upper() == "ON"
RAG_CONTEXTUAL_PREFIX_MAX_TOKENS = int(os.getenv("RAG_CONTEXTUAL_PREFIX_MAX_TOKENS", "100"))
RAG_CONTEXTUAL_DOC_MAX_CHARS     = int(os.getenv("RAG_CONTEXTUAL_DOC_MAX_CHARS", "10000"))
RAG_INGESTION_MODEL     = os.getenv("RAG_INGESTION_MODEL", "").strip()
OPENAI_API_BASE         = os.getenv("OPENAI_API_BASE", "")
OPENAI_MODEL            = os.getenv("OPENAI_MODEL", "")
LOCAL                   = os.getenv("LOCAL", "OFF").strip().upper() == "ON"
OLLAMA_BASE_URL         = os.getenv("OLLAMA_BASE_URL", "")
OLLAMA_MODEL            = os.getenv("OLLAMA_MODEL", "")

SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.markdown', '.rst',
    '.py', '.js', '.jsx', '.ts', '.tsx',
    '.json', '.yaml', '.yml', '.xml',
    '.csv', '.tsv',
    '.pdf', '.docx', '.doc',
    '.html', '.htm',
}

# ── Vérification OCR ──────────────────────────────────────────────────────────
try:
    import pytesseract
    _PYTESSERACT_OK = True
except ImportError:
    _PYTESSERACT_OK = False

try:
    from pdf2image import convert_from_path
    _PDF2IMAGE_OK = True
except ImportError:
    _PDF2IMAGE_OK = False


# ── Schémas ───────────────────────────────────────────────────────────────────

class CollectionCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=120)
    dimension: Optional[int] = None  # None → utilise EMBEDDING_DIMENSION

class IngestRunRequest(BaseModel):
    directory: str = Field(..., description="Chemin absolu du répertoire à indexer")
    collection: str = Field(..., min_length=1)
    recursive: bool = True
    extensions: Optional[List[str]] = None
    use_contextual_chunking: Optional[bool] = None  # None → valeur du .env

class CollectionOut(BaseModel):
    name: str
    vectors_count: int = 0

class SourceOut(BaseModel):
    source: str
    chunks: int

class IngestStatusOut(BaseModel):
    qdrant_available: bool
    qdrant_url: str
    embedding_model: str
    embedding_dimension: int
    embedding_api_base: str
    contextual_chunking_env: bool
    rag_engine_ok: bool
    ocr_available: bool
    ocr_lang: str


# ── Helpers internes ──────────────────────────────────────────────────────────

def _get_qdrant() -> "QdrantClient":
    if not _QDRANT_OK:
        raise HTTPException(status_code=503, detail="qdrant-client non installé sur le serveur.")
    try:
        client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=QDRANT_TIMEOUT)
        client.get_collections()
        return client
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Impossible de joindre Qdrant : {e}")


def _chunk(text: str) -> List[str]:
    if _RAG_ENGINE_OK:
        return _chunk_text(text, max_tokens=256, overlap_tokens=32, hard_max_tokens=512)
    import re
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    chunks, current, length = [], [], 0
    for s in sentences:
        current.append(s)
        length += len(s)
        if length >= 500:
            chunks.append(" ".join(current))
            current = current[-2:] if len(current) > 2 else []
            length = sum(len(x) for x in current)
    if current:
        chunks.append(" ".join(current))
    return chunks or [text]


def _get_embeddings(texts: List[str]) -> List[List[float]]:
    """Génère les embeddings.

    Délègue à rag_engine._get_embeddings() si disponible — utilise ainsi
    la même clé API et la même config (ucfg + Config) que le reste de l'app.
    Fallback : client OpenAI direct avec EMBEDDING_API_BASE / OPENAI_API_KEY.
    """
    if _RAG_ENGINE_OK and _rag_get_embeddings is not None:
        return _rag_get_embeddings(texts)

    # Fallback : client OpenAI direct
    if not _OPENAI_OK:
        raise RuntimeError("openai non installé et rag_engine indisponible")
    from core.config import Config
    api_base = Config.EMBEDDING_API_BASE or EMBEDDING_API_BASE
    api_key  = Config.OPENAI_API_KEY or OPENAI_API_KEY or "none"
    model    = Config.EMBEDDING_MODEL or EMBEDDING_MODEL
    client   = _OpenAI(base_url=api_base, api_key=api_key)
    all_embs = []
    for i in range(0, len(texts), 64):
        batch = texts[i:i + 64]
        resp = client.embeddings.create(input=batch, model=model, encoding_format="float")
        all_embs.extend(item.embedding for item in resp.data)
    return all_embs


def _contextual_prefix_batch_local(document: str, chunks: List[str]) -> List[str]:
    """Fallback autonome pour le chunking contextuel quand rag_engine n'est pas importable.
    Reproduit la logique d'ingest3.py._contextual_prefix_batch_local() sans dépendance
    à l'objet Config — lit directement les variables d'environnement du module.
    """
    if not _OPENAI_OK:
        return [""] * len(chunks)

    doc_excerpt = document[:RAG_CONTEXTUAL_DOC_MAX_CHARS]
    prefixes: List[str] = []

    system_prompt = (
        "Tu es un assistant qui aide à contextualiser des extraits de documents. "
        "Réponds uniquement avec le contexte demandé, sans introduction."
    )

    for i, chunk in enumerate(chunks):
        user_prompt = (
            "<document>\n"
            f"{doc_excerpt}\n"
            "</document>\n\n"
            "Voici l'extrait à contextualiser :\n"
            "<chunk>\n"
            f"{chunk}\n"
            "</chunk>\n\n"
            "Génère une courte phrase (max 2 phrases) qui situe cet extrait dans "
            "le document : de quoi parle-t-il globalement, dans quel contexte "
            "apparaît-il ? Ne répète pas le contenu de l'extrait."
        )
        prefix = ""
        ingestion_model = RAG_INGESTION_MODEL or OPENAI_MODEL

        # API OpenAI-compatible
        if not LOCAL and OPENAI_API_BASE and OPENAI_API_KEY and ingestion_model:
            try:
                llm = _OpenAI(base_url=OPENAI_API_BASE, api_key=OPENAI_API_KEY)
                resp = llm.chat.completions.create(
                    model=ingestion_model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": user_prompt},
                    ],
                    max_tokens=RAG_CONTEXTUAL_PREFIX_MAX_TOKENS,
                    temperature=0.1,
                )
                prefix = resp.choices[0].message.content.strip()
            except Exception as e:
                _log.warning("[ingest_admin] [CTX] chunk %d/%d — échec API : %s", i + 1, len(chunks), e)

        # Ollama
        elif LOCAL and OLLAMA_BASE_URL and OLLAMA_MODEL:
            try:
                import urllib.request as _ur
                import json as _j
                payload = _j.dumps({
                    "model": OLLAMA_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user",   "content": user_prompt},
                    ],
                    "stream": False,
                    "options": {"num_predict": RAG_CONTEXTUAL_PREFIX_MAX_TOKENS, "temperature": 0.1},
                }).encode()
                req = _ur.Request(
                    f"{OLLAMA_BASE_URL.rstrip('/')}/api/chat",
                    data=payload,
                    headers={"Content-Type": "application/json"},
                )
                with _ur.urlopen(req, timeout=20) as r:
                    prefix = _j.loads(r.read())["message"]["content"].strip()
            except Exception as e:
                _log.warning("[ingest_admin] [CTX] chunk %d/%d — échec Ollama : %s", i + 1, len(chunks), e)

        prefixes.append(prefix)

    return prefixes


def _upsert(client: "QdrantClient", collection: str, points: List["PointStruct"]) -> bool:
    import time
    bs = QDRANT_BATCH_SIZE
    for i in range(0, len(points), bs):
        batch = points[i:i + bs]
        for attempt in range(QDRANT_MAX_RETRIES):
            try:
                client.upsert(collection_name=collection, points=batch)
                break
            except Exception as e:
                if attempt == QDRANT_MAX_RETRIES - 1:
                    _log.error("[ingest_admin] upsert failed: %s", e)
                    return False
                time.sleep(2 ** attempt)
    return True


def _extract_text(file_path: Path) -> str:
    """Extraction de texte depuis un fichier (PDF, DOCX, DOC, texte brut)."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = ""
        try:
            import fitz
            doc = fitz.open(str(file_path))
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
        except Exception:
            pass
        if len(text.strip()) < OCR_TEXT_THRESHOLD and _PYTESSERACT_OK and _PDF2IMAGE_OK:
            try:
                images = convert_from_path(str(file_path), dpi=OCR_DPI)
                text = "\n\n".join(pytesseract.image_to_string(img, lang=OCR_LANG) for img in images)
            except Exception as e:
                _log.warning("[ingest_admin] OCR failed for %s: %s", file_path.name, e)
        return text

    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append("  ".join(cells))
            return "\n".join(parts)
        except Exception:
            pass
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(str(file_path), 'r') as z:
                parts = []
                for name in z.namelist():
                    if name.startswith('word/') and name.endswith('.xml') and 'rels' not in name:
                        with z.open(name) as f:
                            root = ET.parse(f).getroot()
                            texts = [n.text for n in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if n.text]
                            if texts:
                                parts.append(" ".join(texts))
                return "\n".join(parts)
        except Exception:
            pass
        return ""

    elif suffix == ".doc":
        # Méthode 1 : antiword
        try:
            import subprocess
            r = subprocess.run(['antiword', str(file_path)], capture_output=True, text=True, timeout=30)
            if r.returncode == 0 and r.stdout:
                return r.stdout
        except Exception:
            pass

        # Méthode 2 : textract
        try:
            import textract
            text = textract.process(str(file_path)).decode('utf-8', errors='replace')
            if text.strip():
                return text
        except Exception:
            pass

        # Méthode 3 : python-docx (récupération partielle)
        try:
            from docx import Document as _Document
            doc = _Document(str(file_path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if text.strip():
                return text
        except Exception:
            pass

        # Méthode 4 : olefile — extraction brute (dernière tentative)
        try:
            import re as _re
            import olefile
            if olefile.isOleFile(str(file_path)):
                ole = olefile.OleFileIO(str(file_path))
                if ole.exists('WordDocument'):
                    data = ole.openstream('WordDocument').read()
                    text = data.decode('latin-1', errors='ignore')
                    text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                    text = _re.sub(
                        r'[^\w\s\.,;:!?\-\'\"àâäéèêëïîôùûüÿçÀÂÄÉÈÊËÏÎÔÙÛÜŸÇ()\[\]{}/@#%&+=]+',
                        '', text
                    )
                    ole.close()
                    if len(text.strip()) > 100:
                        return text
        except Exception:
            pass

        _log.warning("[ingest_admin] Format .doc non extractible : %s", file_path.name)
        return ""

    else:
        try:
            return file_path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""


def _ingest_one_file(
    file_path: Path,
    collection: str,
    qclient: "QdrantClient",
    use_ctx: bool,
    original_name: str | None = None,
) -> int:
    """Ingère un fichier, retourne le nombre de chunks indexés.

    original_name : nom du fichier tel qu'envoyé par l'utilisateur.
        Si fourni, il est utilisé comme valeur du champ ``source`` dans Qdrant
        à la place du nom du fichier temporaire (ex. tmpXXXXXX.pdf).
    """
    text = _extract_text(file_path)
    if not text or len(text.strip()) < 50:
        return 0

    chunks = _chunk(text)
    if not chunks:
        return 0

    context_prefixes: List[str] = [""] * len(chunks)
    if use_ctx:
        if _RAG_ENGINE_OK:
            try:
                context_prefixes = _contextual_prefix_batch(text, chunks)
            except Exception as e:
                _log.warning("[ingest_admin] contextual prefix (rag_engine) failed: %s — fallback local", e)
                context_prefixes = _contextual_prefix_batch_local(text, chunks)
        else:
            context_prefixes = _contextual_prefix_batch_local(text, chunks)

    texts_to_embed = [
        f"{p}\n\n{c}" if p else c
        for p, c in zip(context_prefixes, chunks)
    ]

    embeddings = _get_embeddings(texts_to_embed)
    if not embeddings or len(embeddings) != len(chunks):
        return 0

    source_name = original_name if original_name else file_path.name

    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=emb,
            payload={
                "text": chunk,
                "source": source_name,
                "file_path": str(file_path),
                "conversation_id": "global",
                **({("context_prefix"): prefix} if prefix else {}),
            }
        )
        for chunk, emb, prefix in zip(chunks, embeddings, context_prefixes)
    ]

    return len(chunks) if _upsert(qclient, collection, points) else 0


# ── Routes ────────────────────────────────────────────────────────────────────

@router.get("/ingest/status", response_model=IngestStatusOut)
async def ingest_status(admin: dict = Depends(require_admin)):
    """Retourne l'état de Qdrant et la configuration d'ingestion."""
    available = False
    if _QDRANT_OK:
        try:
            c = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=10)
            c.get_collections()
            available = True
        except Exception:
            pass

    return IngestStatusOut(
        qdrant_available=available,
        qdrant_url=QDRANT_URL,
        embedding_model=EMBEDDING_MODEL,
        embedding_dimension=EMBEDDING_DIMENSION,
        embedding_api_base=EMBEDDING_API_BASE,
        contextual_chunking_env=RAG_CONTEXTUAL_CHUNKING,
        rag_engine_ok=_RAG_ENGINE_OK,
        ocr_available=_PYTESSERACT_OK and _PDF2IMAGE_OK,
        ocr_lang=OCR_LANG,
    )


@router.get("/ingest/collections", response_model=List[CollectionOut])
async def list_collections(admin: dict = Depends(require_admin)):
    """Liste toutes les collections Qdrant avec leur nombre de vecteurs."""
    qclient = _get_qdrant()
    result = []
    for col in qclient.get_collections().collections:
        try:
            info = qclient.get_collection(col.name)
            # points_count est le champ fiable depuis qdrant-client 1.7+
            # vectors_count est déprécié et retourne souvent 0
            count = (
                getattr(info, "points_count", None)
                or getattr(info, "vectors_count", None)
                or 0
            )
        except Exception:
            count = 0
        result.append(CollectionOut(name=col.name, vectors_count=count))
    return result


@router.post("/ingest/collections", response_model=CollectionOut, status_code=201)
async def create_collection(
    payload: CollectionCreate,
    admin: dict = Depends(require_admin),
):
    """Crée une nouvelle collection Qdrant."""
    qclient = _get_qdrant()
    dim = payload.dimension or EMBEDDING_DIMENSION
    existing = {c.name for c in qclient.get_collections().collections}
    if payload.name in existing:
        raise HTTPException(status_code=409, detail=f"La collection '{payload.name}' existe déjà.")
    qclient.create_collection(
        collection_name=payload.name,
        vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
    )
    _log.info("[ingest_admin] Collection créée par %s : %s (dim=%d)", admin["username"], payload.name, dim)
    return CollectionOut(name=payload.name, vectors_count=0)


@router.delete("/ingest/collections/{collection_name}", status_code=204)
async def delete_collection(
    collection_name: str,
    admin: dict = Depends(require_admin),
):
    """Supprime une collection Qdrant et tous ses vecteurs."""
    qclient = _get_qdrant()
    existing = {c.name for c in qclient.get_collections().collections}
    if collection_name not in existing:
        raise HTTPException(status_code=404, detail="Collection introuvable.")
    qclient.delete_collection(collection_name)
    _log.info("[ingest_admin] Collection supprimée par %s : %s", admin["username"], collection_name)


@router.delete("/ingest/collections/{collection_name}/clear", status_code=200)
async def clear_collection(
    collection_name: str,
    admin: dict = Depends(require_admin),
):
    """Vide une collection (supprime tous ses vecteurs) sans la détruire."""
    from qdrant_client.models import Filter
    qclient = _get_qdrant()
    existing = {c.name for c in qclient.get_collections().collections}
    if collection_name not in existing:
        raise HTTPException(status_code=404, detail="Collection introuvable.")
    info = qclient.get_collection(collection_name)
    count_before = info.vectors_count or 0
    qclient.delete(collection_name=collection_name, points_selector=Filter(must=[]))
    _log.info("[ingest_admin] Collection vidée par %s : %s (%d vecteurs)", admin["username"], collection_name, count_before)
    return {"deleted": count_before, "collection": collection_name}


@router.get("/ingest/sources", response_model=List[SourceOut])
async def list_sources(
    collection: str = Query(..., description="Nom de la collection"),
    admin: dict = Depends(require_admin),
):
    """Liste les sources (fichiers) indexées dans une collection."""
    qclient = _get_qdrant()
    from qdrant_client.models import ScrollRequest
    sources: dict[str, int] = {}
    offset = None
    while True:
        result = qclient.scroll(
            collection_name=collection,
            limit=100,
            offset=offset,
            with_payload=["source"],
            with_vectors=False,
        )
        points, next_offset = result
        for p in points:
            src = (p.payload or {}).get("source", "(inconnu)")
            sources[src] = sources.get(src, 0) + 1
        if next_offset is None:
            break
        offset = next_offset
    return [SourceOut(source=s, chunks=c) for s, c in sorted(sources.items())]


@router.delete("/ingest/sources", status_code=204)
async def delete_source(
    collection: str = Query(...),
    source: str = Query(...),
    admin: dict = Depends(require_admin),
):
    """Supprime tous les chunks d'une source dans une collection."""
    from qdrant_client.models import Filter, FieldCondition, MatchValue
    qclient = _get_qdrant()
    qclient.delete(
        collection_name=collection,
        points_selector=Filter(
            must=[FieldCondition(key="source", match=MatchValue(value=source))]
        ),
    )
    _log.info("[ingest_admin] Source supprimée par %s : %s / %s", admin["username"], collection, source)


@router.post("/ingest/file")
async def ingest_single_file(
    file: UploadFile = File(...),
    collection: str = Query(..., description="Nom de la collection cible"),
    use_contextual_chunking: Optional[bool] = Query(None),
    admin: dict = Depends(require_admin),
    user_cfg = Depends(get_current_user_config),
):
    """
    Ingère un fichier uploadé directement dans une collection.
    Réponse synchrone — pour plusieurs fichiers, utiliser /ingest/run.
    """
    if not _QDRANT_OK or not _OPENAI_OK:
        raise HTTPException(status_code=503, detail="Qdrant ou openai non disponible.")

    set_user_config(user_cfg)
    qclient = _get_qdrant()
    use_ctx = use_contextual_chunking if use_contextual_chunking is not None else RAG_CONTEXTUAL_CHUNKING

    suffix = Path(file.filename or "upload").suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = Path(tmp.name)

    try:
        chunks = await asyncio.to_thread(_ingest_one_file, tmp_path, collection, qclient, use_ctx, file.filename)
        _log.info("[ingest_admin] %s ingéré par %s → %d chunks (col=%s)", file.filename, admin["username"], chunks, collection)
        return {"filename": file.filename, "chunks": chunks, "collection": collection}
    except Exception as e:
        _log.exception("[ingest_admin] Erreur ingestion %s", file.filename)
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        tmp_path.unlink(missing_ok=True)


@router.post("/ingest/run")
async def ingest_directory(
    payload: IngestRunRequest,
    admin: dict = Depends(require_admin),
    user_cfg = Depends(get_current_user_config),
):
    """
    Lance l'ingestion d'un répertoire complet avec progression SSE.

    Chaque événement SSE est au format :
        data: {"done": 1, "total": 10, "filename": "doc.pdf", "chunks": 42, "status": "ok"}

    L'événement final est :
        data: {"done": true, "total_chunks": 123, "success": 8, "errors": 2}
    """
    if not _QDRANT_OK or not _OPENAI_OK:
        raise HTTPException(status_code=503, detail="Qdrant ou openai non disponible.")

    set_user_config(user_cfg)
    directory = Path(payload.directory)
    if not directory.exists() or not directory.is_dir():
        raise HTTPException(status_code=400, detail=f"Répertoire introuvable : {payload.directory}")

    extensions = SUPPORTED_EXTENSIONS
    if payload.extensions:
        extensions = {e if e.startswith('.') else f'.{e}' for e in payload.extensions}

    if payload.recursive:
        files = sorted(f for f in directory.rglob('*') if f.is_file() and f.suffix.lower() in extensions)
    else:
        files = sorted(f for f in directory.glob('*') if f.is_file() and f.suffix.lower() in extensions)

    if not files:
        raise HTTPException(status_code=404, detail="Aucun fichier trouvé dans le répertoire.")

    use_ctx = payload.use_contextual_chunking if payload.use_contextual_chunking is not None else RAG_CONTEXTUAL_CHUNKING

    # S'assurer que la collection existe
    qclient = _get_qdrant()
    existing = {c.name for c in qclient.get_collections().collections}
    if payload.collection not in existing:
        qclient.create_collection(
            collection_name=payload.collection,
            vectors_config=VectorParams(size=EMBEDDING_DIMENSION, distance=Distance.COSINE),
        )
        _log.info("[ingest_admin] Collection auto-créée : %s", payload.collection)

    _log.info(
        "[ingest_admin] Ingestion démarrée par %s : %d fichiers → %s (ctx=%s)",
        admin["username"], len(files), payload.collection, use_ctx
    )

    async def _stream():
        total = len(files)
        total_chunks = 0
        success = 0
        errors = 0

        for i, file_path in enumerate(files):
            try:
                chunks = await asyncio.to_thread(
                    _ingest_one_file, file_path, payload.collection, qclient, use_ctx
                )
                if chunks > 0:
                    total_chunks += chunks
                    success += 1
                    event = json.dumps({
                        "done": i + 1, "total": total,
                        "filename": file_path.name,
                        "chunks": chunks,
                        "status": "ok",
                    })
                else:
                    errors += 1
                    event = json.dumps({
                        "done": i + 1, "total": total,
                        "filename": file_path.name,
                        "chunks": 0,
                        "status": "skipped",
                    })
            except Exception as e:
                errors += 1
                event = json.dumps({
                    "done": i + 1, "total": total,
                    "filename": file_path.name,
                    "chunks": 0,
                    "status": "error",
                    "error": str(e),
                })
            yield f"data: {event}\n\n"

        final = json.dumps({
            "done": True,
            "total": total,
            "total_chunks": total_chunks,
            "success": success,
            "errors": errors,
        })
        yield f"data: {final}\n\n"

    return StreamingResponse(_stream(), media_type="text/event-stream")


# ══════════════════════════════════════════════════════════════════════════════
#  Routes personnelles — accessibles à tous les utilisateurs authentifiés
#  Limitées à la collection propre de l'utilisateur (user_cfg.QDRANT_COLLECTION)
# ══════════════════════════════════════════════════════════════════════════════

@router.get("/ingest/personal/status")
async def personal_ingest_status(
    user_cfg = Depends(get_current_user_config),
):
    """Statut Qdrant + nom de la collection personnelle de l'utilisateur."""
    available = False
    if _QDRANT_OK:
        try:
            c = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=10)
            c.get_collections()
            available = True
        except Exception:
            pass

    collection = user_cfg.QDRANT_COLLECTION
    vectors_count = 0
    if available:
        try:
            qclient = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=QDRANT_TIMEOUT)
            info = qclient.get_collection(collection)
            vectors_count = getattr(info, "points_count", None) or getattr(info, "vectors_count", None) or 0
        except Exception:
            pass  # Collection pas encore créée

    return {
        **IngestStatusOut(
            qdrant_available=available,
            qdrant_url=QDRANT_URL,
            embedding_model=EMBEDDING_MODEL,
            embedding_dimension=EMBEDDING_DIMENSION,
            embedding_api_base=EMBEDDING_API_BASE,
            contextual_chunking_env=RAG_CONTEXTUAL_CHUNKING,
            rag_engine_ok=_RAG_ENGINE_OK,
            ocr_available=_PYTESSERACT_OK and _PDF2IMAGE_OK,
            ocr_lang=OCR_LANG,
        ).model_dump(),
        "personal_collection": collection,
        "personal_vectors_count": vectors_count,
    }


@router.get("/ingest/personal/sources", response_model=List[SourceOut])
async def personal_list_sources(
    user_cfg = Depends(get_current_user_config),
):
    """Liste les sources indexées dans la collection personnelle de l'utilisateur."""
    set_user_config(user_cfg)
    collection = user_cfg.QDRANT_COLLECTION
    qclient = _get_qdrant()

    # Créer la collection si elle n'existe pas encore
    existing = {c.name for c in qclient.get_collections().collections}
    if collection not in existing:
        return []

    sources: dict[str, int] = {}
    offset = None
    while True:
        result = qclient.scroll(
            collection_name=collection,
            limit=100,
            offset=offset,
            with_payload=["source"],
            with_vectors=False,
        )
        points, next_offset = result
        for p in points:
            src = (p.payload or {}).get("source", "(inconnu)")
            sources[src] = sources.get(src, 0) + 1
        if next_offset is None:
            break
        offset = next_offset
    return [SourceOut(source=s, chunks=c) for s, c in sorted(sources.items())]


@router.delete("/ingest/personal/sources", status_code=204)
async def personal_delete_source(
    source: str = Query(...),
    user_cfg = Depends(get_current_user_config),
):
    """Supprime une source de la collection personnelle de l'utilisateur."""
    from qdrant_client.models import Filter, FieldCondition, MatchValue
    set_user_config(user_cfg)
    qclient = _get_qdrant()
    qclient.delete(
        collection_name=user_cfg.QDRANT_COLLECTION,
        points_selector=Filter(
            must=[FieldCondition(key="source", match=MatchValue(value=source))]
        ),
    )
    _log.info("[ingest_admin] Source perso supprimée par %s : %s", user_cfg.username, source)


@router.post("/ingest/personal/file")
async def personal_ingest_file(
    file: UploadFile = File(...),
    use_contextual_chunking: Optional[bool] = Query(None),
    user_cfg = Depends(get_current_user_config),
):
    """Ingère un fichier dans la collection personnelle de l'utilisateur."""
    if not _QDRANT_OK or not _OPENAI_OK:
        raise HTTPException(status_code=503, detail="Qdrant ou openai non disponible.")

    set_user_config(user_cfg)
    collection = user_cfg.QDRANT_COLLECTION
    qclient = _get_qdrant()
    use_ctx = use_contextual_chunking if use_contextual_chunking is not None else RAG_CONTEXTUAL_CHUNKING

    # Auto-créer la collection si nécessaire
    existing = {c.name for c in qclient.get_collections().collections}
    if collection not in existing:
        qclient.create_collection(
            collection_name=collection,
            vectors_config=VectorParams(size=EMBEDDING_DIMENSION, distance=Distance.COSINE),
        )
        _log.info("[ingest_admin] Collection personnelle créée : %s", collection)

    suffix = Path(file.filename or "upload").suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = Path(tmp.name)

    try:
        chunks = await asyncio.to_thread(_ingest_one_file, tmp_path, collection, qclient, use_ctx, file.filename)
        _log.info("[ingest_admin] Fichier perso ingéré par %s : %s → %d chunks",
                  user_cfg.username, file.filename, chunks)
        return {"filename": file.filename, "chunks": chunks, "collection": collection}
    except Exception as e:
        _log.exception("[ingest_admin] Erreur ingestion perso %s", file.filename)
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        tmp_path.unlink(missing_ok=True)


@router.delete("/ingest/personal/collection", status_code=200)
async def personal_clear_collection(
    user_cfg = Depends(get_current_user_config),
):
    """Vide intégralement la collection personnelle de l'utilisateur.

    Supprime tous les vecteurs de la collection sans la détruire,
    afin qu'elle reste prête à recevoir de nouveaux documents.
    Retourne le nombre de vecteurs supprimés.
    """
    set_user_config(user_cfg)
    collection = user_cfg.QDRANT_COLLECTION
    qclient = _get_qdrant()

    existing = {c.name for c in qclient.get_collections().collections}
    if collection not in existing:
        return {"deleted": 0, "collection": collection}

    # Compter avant suppression pour retourner un delta significatif
    info = qclient.get_collection(collection)
    count_before = info.vectors_count or 0

    # Supprimer tous les points sans détruire la collection
    from qdrant_client.models import Filter
    qclient.delete(
        collection_name=collection,
        points_selector=Filter(must=[]),  # filtre vide = tous les points
    )

    _log.info(
        "[ingest_admin] Collection personnelle vidée par %s : %s (%d vecteurs supprimés)",
        user_cfg.username, collection, count_before,
    )
    return {"deleted": count_before, "collection": collection}
