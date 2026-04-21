# ============================================================================
# Prométhée — Assistant IA avancé
# ============================================================================
# Auteur  : Pierre COUGET ktulu.analog@gmail.com
# Licence : GNU Affero General Public License v3.0 (AGPL-3.0)
#           https://www.gnu.org/licenses/agpl-3.0.html
# Année   : 2026
# ----------------------------------------------------------------------------
# Ce fichier fait partie du projet Prométhée.
# Vous pouvez le redistribuer et/ou le modifier selon les termes de la
# licence AGPL-3.0 publiée par la Free Software Foundation.
# ============================================================================

"""
tools/docx_reformulation_tools.py — Reformulation de transcriptions .docx
==========================================================================

Problématique :
    Les comptes rendus de réunion transcrits mot à mot (50 à 80 pages)
    sont coûteux pour une transcription one-shot par un LLM. Ce module découpe
    automatiquement le texte en morceaux de taille contrôlée, reformule
    chaque morceau indépendamment, puis assemble le résultat final dans un
    fichier .docx propre.

Outils exposés (2) :

  Outil principal :
    - reformuler_docx        : pipeline complet — lit le .docx source, découpe
                               en chunks, reformule chaque chunk via LLM, ré-assemble
                               et exporte le résultat dans un .docx de sortie.

  Outil auxiliaire (inspection / débogage) :
    - inspecter_docx_source  : analyse un .docx de transcription sans le reformuler —
                               compte les paragraphes, estime le nombre de tokens,
                               propose un plan de découpage.

Stratégie de découpage (chunking)
──────────────────────────────────
  Le texte extrait du .docx est découpé en chunks dont la taille est mesurée
  en tokens estimés (1 token ≈ 4 caractères). Le découpage respecte les
  limites de paragraphes pour ne jamais couper une phrase en deux.

  Le chevauchement (overlap) sert uniquement de contexte de transition : le LLM
  reçoit un marqueur explicite délimitant la zone de contexte (à ne pas reformuler)
  du contenu principal (à reformuler). La sortie ne contient que la partie principale,
  ce qui élimine tout risque de doublon à l'assemblage.

  Paramètres configurables :
    - chunk_tokens   : taille cible d'un chunk (défaut : 8 000 tokens)
    - overlap_tokens : chevauchement entre chunks pour préserver le contexte
                       de transition (défaut : 400 tokens)
    - max_tokens_out : budget de sortie par chunk (défaut : 7 000 tokens)

  Ces valeurs permettent de reformuler un document de 80 pages en ~8-12 appels
  LLM sans jamais approcher la limite de contexte des modèles 128k.

Prompt de reformulation
───────────────────────
  Par défaut, le skill `guide_redacteur` est automatiquement injecté dans le
  prompt système si disponible. Le prompt peut être surchargé via le paramètre
  `instructions_supplementaires`.

  Comportement du LLM attendu :
    - Reformuler en français écrit soigné (registre administratif / professionnel)
    - Supprimer les tics de langage oraux (« euh », répétitions, phrases inachevées)
    - Conserver fidèlement le fond : décisions, noms propres, chiffres, dates
    - Ne pas résumer : la longueur de sortie doit rester proche de l'entrée
    - Structurer en paragraphes courts et lisibles

Sortie
──────
  - Un fichier .docx avec le texte reformulé, stylisé avec une police et
    des marges professionnelles.
  - Un rapport de traitement en JSON : chunks traités, tokens estimés,
    durée, erreurs éventuelles par chunk.

Prérequis :
    pip install python-docx openai
    Configuration via .env : OPENAI_API_BASE, OPENAI_API_KEY, OPENAI_MODEL
    (ou OLLAMA_BASE_URL / OLLAMA_MODEL si LOCAL=ON)
"""

import json
import os
import re
import time
from pathlib import Path
from typing import Optional

from core.tools_engine import tool, set_current_family, _TOOL_ICONS
from core.skill_manager import get_skill_manager
from core.llm_clients import build_family_client
from core.config import Config

set_current_family("docx_reformulation_tools", "Reformulation de documents", "✍️")

_TOOL_ICONS.update({
    "reformuler_docx":       "✍️",
    "inspecter_docx_source": "🔍",
})

# ── Constantes ────────────────────────────────────────────────────────────────

_DEFAULT_CHUNK_TOKENS   = 8_000   # taille cible d'un chunk en tokens estimés
_DEFAULT_OVERLAP_TOKENS = 400     # chevauchement pour préserver le contexte
_DEFAULT_MAX_OUT_TOKENS = 7_000   # budget de réponse par chunk
_CHARS_PER_TOKEN        = 4       # estimation : 1 token ≈ 4 caractères

_EXPORT_DIR = Path.home() / "Exports" / "Prométhée"

# ── Prompt système par défaut ─────────────────────────────────────────────────

_PROMPT_SYSTEME_BASE = """\
Tu es un rédacteur expert chargé de reformuler des transcriptions de réunions.

MISSION :
Reformule le texte fourni en français écrit soigné, de registre professionnel
et administratif. Le texte source est une transcription mot à mot d'échanges oraux.

RÈGLES ABSOLUES :
1. CONSERVER INTÉGRALEMENT le fond : toutes les décisions, tous les chiffres,
   toutes les dates, tous les noms propres, tous les engagements pris.
2. NE PAS RÉSUMER : la longueur du texte reformulé doit rester proche de celle
   de l'original. Ne supprime que les véritables scories orales.
3. SUPPRIMER les scories orales : répétitions, hésitations (« euh », « ben »,
   « voilà »), phrases inachevées, redondances évidentes.
4. STRUCTURER en paragraphes courts et homogènes (4 à 8 lignes).
5. RESPECTER les intervenants : si des prises de parole sont attribuées à des
   personnes (« M. Dupont : … »), conserver l'attribution.
6. NE PAS inventer ni interpréter : en cas d'ambiguïté, reformuler prudemment
   en restant aussi proche que possible de l'intention visible.
7. Retourner UNIQUEMENT le texte reformulé, sans commentaire ni préambule.
"""

# ── Helpers internes ──────────────────────────────────────────────────────────

def _ok(data: dict) -> str:
    return json.dumps({"status": "ok", **data}, ensure_ascii=False)

def _err(msg: str) -> str:
    return json.dumps({"status": "error", "error": msg}, ensure_ascii=False)

def _est_tokens(text: str) -> int:
    """Estimation rapide du nombre de tokens (1 token ≈ 4 caractères)."""
    return max(1, len(text) // _CHARS_PER_TOKEN)

def _resolve_output(output_path: str, default_name: str) -> Path:
    if output_path:
        p = Path(output_path).expanduser()
        if not p.is_absolute():
            p = Path.home() / p
    else:
        _EXPORT_DIR.mkdir(parents=True, exist_ok=True)
        p = _EXPORT_DIR / default_name
    p.parent.mkdir(parents=True, exist_ok=True)
    return p

def _extraire_texte_docx(docx_path: Path) -> list[str]:
    """
    Extrait les paragraphes d'un .docx sous forme de liste de chaînes.
    Les paragraphes vides sont conservés comme séparateurs sémantiques.
    """
    from docx import Document
    doc = Document(str(docx_path))
    paragraphes = [p.text for p in doc.paragraphs]
    return paragraphes

def _charger_skill_redacteur() -> str:
    """Charge le skill guide_redacteur si disponible, sinon retourne ''."""
    try:
        sm = get_skill_manager()
        return sm.read_skill("guide_redacteur")
    except Exception:
        return ""

def _construire_prompt_systeme(instructions_supplementaires: str = "") -> str:
    """
    Construit le prompt système en combinant :
    - le prompt de base intégré à ce fichier
    - le contenu du skill guide_redacteur (si disponible)
    - les instructions supplémentaires passées par l'utilisateur
    """
    skill = _charger_skill_redacteur()
    parties = [_PROMPT_SYSTEME_BASE]
    if skill:
        parties.append("\n\n── GUIDE RÉDACTIONNEL (skill guide_redacteur) ──\n\n" + skill)
    if instructions_supplementaires and instructions_supplementaires.strip():
        parties.append(
            "\n\n── INSTRUCTIONS SPÉCIFIQUES POUR CE DOCUMENT ──\n\n"
            + instructions_supplementaires.strip()
        )
    return "\n".join(parties)

def _decouper_en_chunks(
    paragraphes: list[str],
    chunk_tokens: int = _DEFAULT_CHUNK_TOKENS,
    overlap_tokens: int = _DEFAULT_OVERLAP_TOKENS,
) -> list[dict]:
    """
    Découpe la liste de paragraphes en chunks de taille contrôlée.

    Chaque chunk est un dict :
      {
        "index":        int,   # numéro du chunk (0-based)
        "contexte":     str,   # overlap du chunk précédent (contexte de transition,
                               #   à ne PAS reformuler — vide pour le premier chunk)
        "texte":        str,   # texte principal à reformuler
        "tokens_est":   int,   # tokens estimés (contexte + texte)
        "para_debut":   int,   # indice du premier paragraphe du texte principal
        "para_fin":     int,   # indice du dernier paragraphe du texte principal
      }

    L'algorithme :
    - Accumule les paragraphes jusqu'à atteindre chunk_tokens.
    - Calcule un chevauchement (overlap) de overlap_tokens depuis la fin du
      chunk courant, qui sera passé comme contexte de transition au chunk suivant.
    - Stocke le contexte séparément du texte principal : le LLM reçoit une
      instruction explicite indiquant que la zone contexte est fournie pour
      la cohérence stylistique uniquement et ne doit PAS être reformulée.
    - Ne coupe jamais un paragraphe en deux.
    """
    chunks: list[dict] = []
    n = len(paragraphes)
    if n == 0:
        return chunks

    chunk_idx    = 0
    para_start   = 0
    contexte_buf = ""   # overlap issu du chunk précédent (contexte seul, non reformulé)

    while para_start < n:
        # ── Accumulation du texte principal ──────────────────────────────────
        tokens_acc  = 0
        texte_chunk = ""
        para_i      = para_start

        while para_i < n and tokens_acc < chunk_tokens:
            p = paragraphes[para_i]
            tokens_acc += _est_tokens(p) + 1  # +1 pour le saut de ligne
            texte_chunk += ("\n" if texte_chunk else "") + p
            para_i += 1

        # ── Calcul du chevauchement pour le chunk suivant ─────────────────────
        # On prend les derniers paragraphes du texte courant dont la somme ≤ overlap_tokens.
        next_contexte = ""
        if overlap_tokens > 0:
            overlap_toks = 0
            para_ov      = para_i - 1
            while para_ov >= para_start and overlap_toks < overlap_tokens:
                p = paragraphes[para_ov]
                overlap_toks += _est_tokens(p) + 1
                para_ov -= 1
            overlap_start = para_ov + 1
            if overlap_start < para_i:
                next_contexte = "\n".join(paragraphes[overlap_start:para_i])

        chunks.append({
            "index":      chunk_idx,
            "contexte":   contexte_buf,
            "texte":      texte_chunk.strip(),
            "tokens_est": _est_tokens(contexte_buf) + _est_tokens(texte_chunk),
            "para_debut": para_start,
            "para_fin":   para_i - 1,
        })

        chunk_idx    += 1
        para_start    = para_i
        contexte_buf  = next_contexte

    return chunks

def _reformuler_via_llm(
    texte: str,
    prompt_systeme: str,
    max_tokens_out: int,
    modele: str | None = None,
    contexte: str = "",
) -> tuple[str, int]:
    """
    Envoie un chunk au LLM et retourne (texte_reformulé, tokens_utilisés).
    Utilise build_family_client() pour respecter la configuration Albert du .env.

    Si `contexte` est fourni (overlap du chunk précédent), il est injecté avant
    le texte principal avec un marqueur explicite indiquant au LLM de NE PAS le
    reformuler — il sert uniquement à préserver la cohérence stylistique de
    transition. Seul le contenu délimité par les marqueurs DÉBUT/FIN doit être
    retourné, ce qui élimine tout risque de doublon à l'assemblage.

    Lève une exception en cas d'erreur API.
    """
    client, resolved_model = build_family_client("docx_reformulation_tools")
    model_to_use = modele or resolved_model

    if contexte and contexte.strip():
        contenu_user = (
            "── CONTEXTE DE TRANSITION (NE PAS REFORMULER) ──\n"
            "Le passage ci-dessous provient de la fin du bloc précédent. "
            "Il t'est fourni uniquement pour que ta reformulation s'y enchaîne "
            "de façon naturelle. Ne le réécris PAS, ne le répète PAS dans ta réponse.\n\n"
            + contexte.strip()
            + "\n\n── FIN DU CONTEXTE ──\n\n"
            "── TEXTE À REFORMULER — DÉBUT ──\n"
            + texte
            + "\n── TEXTE À REFORMULER — FIN ──\n\n"
            "Retourne UNIQUEMENT la reformulation du texte délimité ci-dessus, "
            "sans commentaire ni répétition du contexte."
        )
    else:
        contenu_user = (
            "Voici le texte à reformuler. Retourne uniquement le texte "
            "reformulé, sans commentaire.\n\n"
            "─────────────────────────────────────\n"
            + texte
            + "\n─────────────────────────────────────"
        )

    resp = client.chat.completions.create(
        model=model_to_use,
        max_tokens=max_tokens_out,
        messages=[
            {"role": "system", "content": prompt_systeme},
            {"role": "user",   "content": contenu_user},
        ],
    )
    texte_out   = resp.choices[0].message.content or "" if resp.choices else ""
    usage       = resp.usage
    tokens_used = (getattr(usage, "prompt_tokens", 0) or 0) + (getattr(usage, "completion_tokens", 0) or 0)
    return texte_out, tokens_used

def _assembler_docx(
    sections_reformulees: list[str],
    titre_document: str,
    output_path: Path,
) -> None:
    """
    Assemble les sections reformulées dans un fichier .docx propre.
    Applique un style professionnel : police Calibri, interligne 1.15,
    marges standards.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Style Normal personnalisé
    style_normal = doc.styles["Normal"]
    font = style_normal.font
    font.name      = "Calibri"
    font.size      = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style_normal.paragraph_format
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(14)

    # Titre du document
    if titre_document:
        titre_para = doc.add_heading(titre_document, level=1)
        titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titre_para.paragraph_format.space_after = Pt(18)
        for run in titre_para.runs:
            run.font.name  = "Calibri"
            run.font.size  = Pt(16)
            run.font.color.rgb = RGBColor(0x22, 0x55, 0xA4)

    # Contenu reformulé
    for i, section_text in enumerate(sections_reformulees):
        if not section_text or not section_text.strip():
            continue

        # Chaque ligne non vide devient un paragraphe
        lignes = [l.strip() for l in section_text.split("\n")]
        for ligne in lignes:
            if ligne:
                p = doc.add_paragraph(ligne)
                p.style = doc.styles["Normal"]
            # ligne vide → espace supplémentaire (ne pas créer un paragraphe vide)

        # Séparateur visuel léger entre chunks (sauf le dernier)
        if i < len(sections_reformulees) - 1:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after  = Pt(4)

    doc.save(str(output_path))


# ═══════════════════════════════════════════════════════════════════════════════
# OUTIL PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

@tool(
    name="reformuler_docx",
    description=(
        "Reformule une transcription de réunion au format .docx (forme orale, mot à mot) "
        "en français écrit soigné et professionnel. "
        "Conçu pour des documents très longs (50 à 80 pages) : le texte est automatiquement "
        "découpé en morceaux de taille contrôlée pour ne jamais saturer la fenêtre de contexte "
        "du LLM. Chaque morceau est reformulé indépendamment, puis les résultats sont assemblés "
        "dans un fichier .docx de sortie. "
        "CORRECTION BUG 3 — RÈGLE IMPÉRATIVE : cet outil DOIT être appelé pour toute "
        "reformulation de .docx. Ne jamais prétendre reformuler ou enregistrer un fichier "
        "sans appeler cet outil — une réponse textuelle sans appel d'outil est une erreur. "
        "RÈGLE IMPORTANTE : appeler skill_read('guide_redacteur') avant cet outil pour "
        "injecter les conventions rédactionnelles spécifiques à l'organisation. "
        "Retourne un rapport détaillé : nombre de chunks traités, tokens consommés, durée."
    ),
    parameters={
        "type": "object",
        "properties": {
            "input_path": {
                "type": "string",
                "description": (
                    "Chemin absolu ou relatif au home du fichier .docx source (transcription). "
                    "Ex : ~/Documents/CR_reunion_2026_01_15.docx"
                ),
            },
            "output_path": {
                "type": "string",
                "description": (
                    "Chemin de destination du .docx reformulé. "
                    "Si omis, créé dans ~/Exports/Prométhée/ avec le suffixe _reformule.docx."
                ),
            },
            "titre_document": {
                "type": "string",
                "description": (
                    "Titre à afficher en tête du document reformulé. "
                    "Si omis, utilise le nom du fichier source."
                ),
            },
            "instructions_supplementaires": {
                "type": "string",
                "description": (
                    "Instructions rédactionnelles spécifiques à ce document : "
                    "contexte de la réunion, noms des intervenants, termes techniques "
                    "à conserver, niveau de détail souhaité, etc. "
                    "Ces instructions sont ajoutées au prompt système."
                ),
            },
            "chunk_tokens": {
                "type": "integer",
                "description": (
                    "Taille cible d'un chunk en tokens estimés (1 token ≈ 4 caractères). "
                    "Défaut : 8000. Réduire si le LLM renvoie des erreurs de contexte. "
                    "Augmenter (max 12000) pour obtenir une reformulation plus cohérente "
                    "sur de grandes plages de texte (modèles 128k)."
                ),
            },
            "overlap_tokens": {
                "type": "integer",
                "description": (
                    "Taille du contexte de transition entre chunks consécutifs, en tokens estimés. "
                    "Ce contexte est transmis au LLM avec un marqueur explicite lui indiquant "
                    "de ne pas le reformuler — il sert uniquement à assurer la cohérence stylistique. "
                    "Défaut : 400. Ne pas dépasser 600."
                ),
            },
            "max_tokens_out": {
                "type": "integer",
                "description": (
                    "Budget de tokens de sortie par chunk (paramètre max_tokens de l'API). "
                    "Défaut : 7000. Augmenter si les chunks semblent tronqués en sortie."
                ),
            },
            "modele": {
                "type": "string",
                "description": (
                    "Modèle LLM à utiliser. Si omis, utilise le modèle assigné "
                    "à la famille d'outils (paramètres) ou Config.active_model(). "
                    "Ne modifier que si nécessaire."
                ),
            },
        },
        "required": ["input_path"],
    },
)
def reformuler_docx(
    input_path: str,
    output_path: str = "",
    titre_document: str = "",
    instructions_supplementaires: str = "",
    chunk_tokens: int = _DEFAULT_CHUNK_TOKENS,
    overlap_tokens: int = _DEFAULT_OVERLAP_TOKENS,
    max_tokens_out: int = _DEFAULT_MAX_OUT_TOKENS,
    modele: str = "",
) -> str:
    try:
        # ── Vérification de la configuration ────────────────────────────────
        if not Config.LOCAL and not Config.OPENAI_API_KEY:
            return _err(
                "reformuler_docx : variable OPENAI_API_KEY absente. "
                "Ajoutez-la dans votre fichier .env."
            )

        # ── Fichier source ───────────────────────────────────────────────────
        src = Path(input_path).expanduser().resolve()
        if not src.exists():
            return _err(f"reformuler_docx : fichier source introuvable : {src}")
        if src.suffix.lower() != ".docx":
            return _err(
                f"reformuler_docx : extension non supportée ({src.suffix}). "
                "Seul le format .docx est accepté."
            )

        # ── Fichier de sortie ────────────────────────────────────────────────
        nom_sortie = src.stem + "_reformule.docx"
        dest = _resolve_output(output_path, nom_sortie)

        # ── Titre du document ────────────────────────────────────────────────
        titre = titre_document.strip() if titre_document else src.stem.replace("_", " ")

        # ── Extraction du texte ──────────────────────────────────────────────
        paragraphes = _extraire_texte_docx(src)
        total_chars = sum(len(p) for p in paragraphes)
        total_tokens_est = _est_tokens("".join(paragraphes))

        # ── Découpage en chunks ──────────────────────────────────────────────
        chunk_tokens   = max(500,   min(chunk_tokens,  12_000))
        overlap_tokens = max(0,     min(overlap_tokens,   600))
        max_tokens_out = max(1000,  min(max_tokens_out, 8_192))

        chunks = _decouper_en_chunks(paragraphes, chunk_tokens, overlap_tokens)

        # ── Prompt système ───────────────────────────────────────────────────
        prompt_sys = _construire_prompt_systeme(instructions_supplementaires)

        # ── Traitement chunk par chunk ────────────────────────────────────────
        sections_reformulees: list[str] = []
        rapport_chunks: list[dict]      = []
        total_tokens_api = 0
        erreurs: list[str]              = []
        t_debut = time.time()

        for chunk in chunks:
            t_chunk = time.time()
            chunk_num = chunk["index"] + 1
            n_chunks  = len(chunks)

            try:
                texte_ref, tokens_used = _reformuler_via_llm(
                    texte          = chunk["texte"],
                    prompt_systeme = prompt_sys,
                    max_tokens_out = max_tokens_out,
                    modele         = modele or None,
                    contexte       = chunk["contexte"],
                )
                sections_reformulees.append(texte_ref)
                total_tokens_api += tokens_used
                rapport_chunks.append({
                    "chunk":       chunk_num,
                    "total":       n_chunks,
                    "tokens_in_est": chunk["tokens_est"],
                    "tokens_api":  tokens_used,
                    "duree_s":     round(time.time() - t_chunk, 1),
                    "statut":      "ok",
                })
            except Exception as e_chunk:
                msg_err = f"Chunk {chunk_num}/{n_chunks} : {e_chunk}"
                erreurs.append(msg_err)
                # On insère le texte original non reformulé en fallback
                sections_reformulees.append(
                    f"[ERREUR DE REFORMULATION — TEXTE ORIGINAL]\n{chunk['texte']}"
                )
                rapport_chunks.append({
                    "chunk":   chunk_num,
                    "total":   n_chunks,
                    "statut":  "erreur",
                    "detail":  str(e_chunk),
                })

        duree_totale = round(time.time() - t_debut, 1)

        # ── Assemblage du .docx de sortie ────────────────────────────────────
        _assembler_docx(sections_reformulees, titre, dest)

        # ── Résultat ─────────────────────────────────────────────────────────
        result = {
            "path":               str(dest),
            "size_bytes":         dest.stat().st_size,
            "fichier_source":     str(src),
            "titre":              titre,
            "paragraphes_source": len(paragraphes),
            "chars_source":       total_chars,
            "tokens_source_est":  total_tokens_est,
            "chunks_traites":     len(chunks),
            "chunks_ok":          sum(1 for c in rapport_chunks if c["statut"] == "ok"),
            "chunks_erreur":      len(erreurs),
            "tokens_api_total":   total_tokens_api,
            "duree_totale_s":     duree_totale,
            "modele":             modele or Config.active_model(),
            "detail_chunks":      rapport_chunks,
        }
        if erreurs:
            result["erreurs"] = erreurs

        return _ok(result)

    except Exception as e:
        return _err(f"reformuler_docx : {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# OUTIL AUXILIAIRE 1 — Inspection
# ═══════════════════════════════════════════════════════════════════════════════

@tool(
    name="inspecter_docx_source",
    description=(
        "Analyse un fichier .docx de transcription sans le reformuler. "
        "Retourne : nombre de paragraphes, nombre de mots, estimation de tokens, "
        "et plan de découpage prévu (nombre de chunks, taille de chaque chunk) "
        "selon les paramètres de chunking choisis. "
        "Utile pour estimer la durée et le coût API avant de lancer reformuler_docx."
    ),
    parameters={
        "type": "object",
        "properties": {
            "input_path": {
                "type": "string",
                "description": "Chemin du fichier .docx à analyser.",
            },
            "chunk_tokens": {
                "type": "integer",
                "description": f"Taille cible d'un chunk en tokens. Défaut : {_DEFAULT_CHUNK_TOKENS}.",
            },
            "overlap_tokens": {
                "type": "integer",
                "description": f"Chevauchement en tokens. Défaut : {_DEFAULT_OVERLAP_TOKENS}.",
            },
        },
        "required": ["input_path"],
    },
)
def inspecter_docx_source(
    input_path: str,
    chunk_tokens: int = _DEFAULT_CHUNK_TOKENS,
    overlap_tokens: int = _DEFAULT_OVERLAP_TOKENS,
) -> str:
    try:
        src = Path(input_path).expanduser().resolve()
        if not src.exists():
            return _err(f"inspecter_docx_source : fichier introuvable : {src}")

        paragraphes = _extraire_texte_docx(src)
        texte_total = "\n".join(paragraphes)
        nb_mots     = len(texte_total.split())
        nb_chars    = len(texte_total)
        nb_tokens   = _est_tokens(texte_total)

        chunks = _decouper_en_chunks(paragraphes, chunk_tokens, overlap_tokens)

        plan_chunks = [
            {
                "index":      c["index"],
                "tokens_est": c["tokens_est"],
                "chars":      len(c["texte"]),
                "para_debut": c["para_debut"],
                "para_fin":   c["para_fin"],
                "extrait":    c["texte"][:120].replace("\n", " ") + "…",
            }
            for c in chunks
        ]

        # Estimation du coût API (à la louche sur base Mistral large en l'absence de données retournées par Albert)
        # input: ~2€/Mtok, output: ~6€/Mtok
        tokens_input_total  = sum(c["tokens_est"] for c in chunks)
        tokens_output_total = len(chunks) * _DEFAULT_MAX_OUT_TOKENS  # majorant
        cout_estime_usd = (tokens_input_total * 2 + tokens_output_total * 6) / 1_000_000

        return _ok({
            "fichier":             str(src),
            "taille_fichier_kb":   round(src.stat().st_size / 1024, 1),
            "nb_paragraphes":      len(paragraphes),
            "nb_mots":             nb_mots,
            "nb_chars":            nb_chars,
            "tokens_est_total":    nb_tokens,
            "nb_chunks_prevus":    len(chunks),
            "chunk_tokens_cible":  chunk_tokens,
            "overlap_tokens":      overlap_tokens,
            "cout_estime_usd_max": round(cout_estime_usd, 3),
            "plan_decoupage":      plan_chunks,
        })

    except Exception as e:
        return _err(f"inspecter_docx_source : {e}")
